import io
import os
import json
import base64
import urllib.parse
from datetime import datetime
from typing import Optional, List, Union

import boto3
import docx
import openpyxl
import pytz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from sqlalchemy.orm import Session
from sqlalchemy import text

from app import models
from app.database import get_db
from auth import get_current_user

# Try safe import for embedding ingestion
try:
    from embedding_service import ingest_document_vector
except ImportError:
    try:
        from app.embedding_service import ingest_document_vector
    except ImportError:
        def ingest_document_vector(*args, **kwargs):
            pass

router = APIRouter(prefix="/api/v1", tags=["Document Upload & Archive"])

# AWS S3 Client Configuration
AWS_REGION = os.getenv("AWS_REGION", "eu-central-1")
BUCKET_NAME = os.getenv("AWS_BUCKET_NAME", "kmp-centralised-security-storage")

s3_client = boto3.client(
    "s3",
    aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
    region_name=AWS_REGION
)

def get_eat_now():
    eat_tz = pytz.timezone("Africa/Nairobi")
    return datetime.now(eat_tz).replace(tzinfo=None)

def get_doc_archive_model():
    """Dynamically resolves the Document Archive model across table name naming conventions."""
    for model_name in ['DocumentArchive', 'Document_Archive', 'document_archive', 'document_archives']:
        if hasattr(models, model_name):
            return getattr(models, model_name)
    raise HTTPException(status_code=500, detail="Document Archive database model is not configured.")

def get_template_model():
    for model_name in ['CommandTemplate', 'Command_Template', 'command_template', 'command_templates']:
        if hasattr(models, model_name):
            return getattr(models, model_name)
    return None

def get_general_doc_model():
    for name in ['GeneralDocuments', 'General_Documents', 'general_documents']:
        if hasattr(models, name):
            return getattr(models, name)
    return None

def log_semantic_audit(db: Session, fnum: str, action: str, target_identifier: str, changes: dict, remarks: str = ""):
    try:
        eat_time = get_eat_now()
        formatted_details = f"Target: {target_identifier} | Changes: " + ", ".join(
            [f"{k}: {v[0]} -> {v[1]}" for k, v in changes.items()]
        ) + f" | Remarks: {remarks}"
        
        audit_model = getattr(models, 'Audit_Logs', getattr(models, 'AuditLogs', None))
        if audit_model:
            new_audit = audit_model(
                event_type=action,
                target_user=target_identifier,
                status="SUCCESS",
                details=formatted_details,
                user_fnum=fnum,
                created_at=eat_time.strftime('%Y-%m-%d %H:%M:%S')
            )
            db.add(new_audit)
            db.commit()
    except Exception as e:
        print(f"Audit Log Notice: {e}")
        db.rollback()

@router.get("/reports/archive")
def get_document_archive(db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    try:
        ArchiveModel = get_doc_archive_model()
        
        docs = db.query(ArchiveModel).all()
        if not docs:
            raw_result = db.execute(text("SELECT id, file_name, doc_type, file_size, file_path, region, station, uploaded_by, upload_date FROM document_archive ORDER BY id DESC")).fetchall()
            return [
                {
                    "id": row[0],
                    "name": row[1] or "document",
                    "type": row[2] or "General Document",
                    "date": str(row[8]).split(' ')[0] if row[8] else "",
                    "size": row[3] or "N/A",
                    "file_path": row[4] or "",
                    "region": row[5] or "KMP HEADQUARTERS",
                    "station": row[6] or "KMP HEADQUARTERS"
                } for row in raw_result
            ]

        results = []
        for doc in docs:
            file_name = getattr(doc, 'file_name', getattr(doc, 'filename', 'document'))
            doc_type = getattr(doc, 'doc_type', getattr(doc, 'doctype', 'General Document'))
            file_size = getattr(doc, 'file_size', getattr(doc, 'filesize', 'N/A'))
            file_path = getattr(doc, 'file_path', getattr(doc, 'filepath', ''))
            region = getattr(doc, 'region', 'KMP HEADQUARTERS')
            station = getattr(doc, 'station', 'KMP HEADQUARTERS')
            upload_date = getattr(doc, 'upload_date', getattr(doc, 'uploaded_at', None))
            
            date_str = ""
            if isinstance(upload_date, datetime):
                date_str = upload_date.strftime("%Y-%m-%d")
            elif upload_date:
                date_str = str(upload_date).split(' ')[0]

            results.append({
                "id": getattr(doc, 'id', getattr(doc, 'sn', 1)),
                "name": file_name,
                "type": doc_type or "General Document",
                "date": date_str,
                "size": file_size or "N/A",
                "file_path": file_path,
                "region": region,
                "station": station
            })
        return results
    except Exception as e:
        print(f"Archive fetch error: {str(e)}")
        try:
            raw_result = db.execute(text("SELECT id, file_name, doc_type, file_size, file_path, region, station, upload_date FROM document_archive ORDER BY id DESC")).fetchall()
            return [
                {
                    "id": row[0], "name": row[1] or "document", "type": row[2] or "General Document",
                    "date": str(row[7]).split(' ')[0] if row[7] else "", "size": row[3] or "N/A",
                    "file_path": row[4] or "", "region": row[5] or "KMP HEADQUARTERS", "station": row[6] or "KMP HEADQUARTERS"
                } for row in raw_result
            ]
        except Exception as inner_err:
            print(f"Raw SQL fallback failed: {inner_err}")
            return []


@router.post("/reports/upload-word-report")
async def upload_word_report(
    file: Optional[UploadFile] = File(None),
    files: Optional[List[UploadFile]] = File(None),
    doc_type: str = Form(...),  
    target_region: Optional[str] = Form(None), 
    target_station: Optional[str] = Form(None), 
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    ArchiveModel = get_doc_archive_model()
    file_list = [f for f in [file, *(files or [])] if f is not None]
    
    if not file_list:
        raise HTTPException(status_code=400, detail="No files received for intake.")

    effective_region = target_region.upper() if (target_region and current_user.role in ["SUPER_ADMIN", "ADMIN"]) else (current_user.region or "KMP GENERAL")
    effective_station = target_station.upper() if (target_station and current_user.role in ["SUPER_ADMIN", "ADMIN"]) else (current_user.station or "KMP HEADQUARTERS")
    eat_time = get_eat_now()
    uploaded_count = 0

    try:
        for single_file in file_list:
            contents = await single_file.read()
            
            if doc_type == "weekly_report" and single_file.filename.lower().endswith('.docx'):
                try:
                    doc = Document(io.BytesIO(contents))
                    for para in doc.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for run in para.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                    formatted_io = io.BytesIO()
                    doc.save(formatted_io)
                    contents = formatted_io.getvalue()
                except Exception as format_err:
                    print(f"Document justification notice: {format_err}")

            file_size_kb = max(1, round(len(contents) / 1024))
            file_size_str = f"{file_size_kb} KB" if file_size_kb < 1024 else f"{round(file_size_kb / 1024, 1)} MB"

            timestamp = eat_time.strftime("%Y%m%d_%H%M%S")
            safe_filename = f"{timestamp}_{single_file.filename.replace(' ', '_')}"
            s3_key = f"reports_archive/{safe_filename}"
            
            s3_client.put_object(
                Bucket=BUCKET_NAME, Key=s3_key, Body=contents,
                ContentType=single_file.content_type or "application/octet-stream", ServerSideEncryption="AES256"
            )
            
            full_s3_url = f"https://{BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{s3_key}"
            display_type = "Weekly Report" if doc_type == "weekly_report" else ("General Document" if doc_type == "general_doc" else doc_type)
            
            new_archive = ArchiveModel(
                file_name=single_file.filename,
                doc_type=display_type,
                file_size=file_size_str,
                file_path=full_s3_url, 
                region=effective_region, 
                station=effective_station, 
                uploaded_by=current_user.fnum,
                upload_date=eat_time 
            )
            db.add(new_archive)
            db.flush() # Flush to get the ID for vector ingestion if needed

            # Safely trigger vector embedding ingestion if available
            try:
                ingest_document_vector(
                    db=db,
                    document_id=new_archive.id,
                    document_type=display_type,
                    title=single_file.filename,
                    raw_text=contents.decode('utf-8', errors='ignore') if single_file.filename.endswith('.txt') else f"Document: {single_file.filename}",
                    region=effective_region,
                    division=getattr(current_user, 'division', 'KMP HEADQUARTERS'),
                    station=effective_station,
                    sd_ref="N/A"
                )
            except Exception as vec_err:
                print(f"Vector ingestion notice: {vec_err}")

            uploaded_count += 1

        db.commit()
        return {"status": "success", "message": f"Successfully archived {uploaded_count} document(s)."}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to process document intake: {str(e)}")

@router.get("/reports/download/{doc_id}")
@router.get("/templates/download/{doc_id}")
def download_archive_file(
    doc_id: int, 
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    ArchiveModel = get_doc_archive_model()
    TemplateModel = get_template_model()
    GeneralDocModel = get_general_doc_model()
    
    # 🟢 1. Search across Archive, Templates, AND General Documents
    doc_record = db.query(ArchiveModel).filter(ArchiveModel.id == doc_id).first()
    if not doc_record and TemplateModel:
        doc_record = db.query(TemplateModel).filter(TemplateModel.id == doc_id).first()
    if not doc_record and GeneralDocModel:
        doc_record = db.query(GeneralDocModel).filter(GeneralDocModel.id == doc_id).first()

    if not doc_record:
        raise HTTPException(status_code=404, detail="Document record not found in system database.")
        
    file_path = getattr(doc_record, 'file_path', getattr(doc_record, 'url', ''))
    file_name = getattr(doc_record, 'file_name', getattr(doc_record, 'name', 'document'))

    if not str(file_path).startswith("http"):
        raise HTTPException(status_code=404, detail="Local static files cannot be dynamically watermarked. S3 storage required.")

    parsed_url = urllib.parse.urlparse(file_path)
    original_s3_key = parsed_url.path.lstrip('/') 
    file_extension = file_name.lower().split('.')[-1] if '.' in file_name else "bin"

    try:
        file_stream = io.BytesIO()
        s3_client.download_fileobj(BUCKET_NAME, original_s3_key, file_stream)
        file_stream.seek(0)
        raw_bytes = file_stream.getvalue()

        eat_time = get_eat_now()
        timestamp_eat = eat_time.strftime("%Y-%m-%d %H:%M:%S EAT")

        officer_fnum = (current_user.fnum or "HQ-UNKNOWN").strip().upper()
        officer_rank = (current_user.rank or "OFFICER").strip().upper()
        officer_name = (current_user.name or "UNKNOWN").strip().upper()
        officer_signature = f"{officer_fnum} {officer_rank} {officer_name}"
        command_post = f"{current_user.station or 'KMP HEADQUARTERS'}, {current_user.region or 'KMP HEADQUARTERS'}"
        stamp_id = f"KMP-STAMP-{officer_fnum}-{eat_time.strftime('%Y%m%d%H%M%S')}"

        # 🟢 2. Compact payload to stay well under the 255-character XML metadata limit
        compact_payload = {"f": officer_fnum, "s": stamp_id}
        encoded_token = base64.b64encode(json.dumps(compact_payload).encode('utf-8')).decode('utf-8')
        keywords_str = f"KMP_AUDIT;{encoded_token}"[:250]
        comments_str = f"Export: {officer_signature} [{command_post}]. ID: {stamp_id}"

        receipt_text = (
            "========================================================\n"
            "         KAMPALA METROPOLITAN POLICE HEADQUARTERS         \n"
            "         SECURE DOCUMENT & TEMPLATES ACCESS         \n"
            "--------------------------------------------------------\n"
            f"ACCESSED BY    : {officer_signature}\n"
            f"CLEARANCE      : {current_user.role} | STATION: {current_user.station}\n"
            f"AUDIT STAMP ID : {stamp_id}\n"
            f"TIMESTAMP      : {timestamp_eat}\n"
            "========================================================"
        )

        output_stream = io.BytesIO()
        content_type = "application/octet-stream"

        if file_extension == 'docx':
            word_doc = Document(io.BytesIO(raw_bytes))
            
            core_props = word_doc.core_properties
            core_props.author = officer_signature
            core_props.last_modified_by = officer_signature
            core_props.keywords = keywords_str
            core_props.comments = comments_str
            core_props.category = "RESTRICTED / LAW ENFORCEMENT RECORD"

            section = word_doc.sections[0]
            footer = section.footer
            footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = footer_p.add_run(receipt_text)
            run.font.name = 'Courier New' 
            run.font.size = Pt(7.5) 
            run.font.bold = True
            run.font.color.rgb = RGBColor(139, 0, 0) 
            word_doc.save(output_stream)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        elif file_extension in ['xlsx', 'xls']:
            wb = openpyxl.load_workbook(io.BytesIO(raw_bytes))
            
            wb.properties.creator = officer_signature
            wb.properties.lastModifiedBy = officer_signature
            wb.properties.keywords = keywords_str
            wb.properties.description = comments_str
            wb.properties.category = "RESTRICTED / FORENSIC POLICE RECORD"

            for ws in wb.worksheets:
                if hasattr(ws, 'sheet_footer'): ws.sheet_footer.center.text = receipt_text
                elif hasattr(ws, 'odd_footer'): ws.odd_footer.center.text = receipt_text
            wb.save(output_stream)
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        else:
            output_stream.write(raw_bytes)

        output_stream.seek(0)

        stamped_s3_key = f"forensic_cache/{officer_fnum}_DOC_{doc_id}.{file_extension}"
        s3_client.upload_fileobj(output_stream, BUCKET_NAME, stamped_s3_key, ExtraArgs={"ContentType": content_type})

        stamped_url = f"https://{BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{stamped_s3_key}"
        return {"download_url": stamped_url, "file_url": stamped_url}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Universal Forensic Stamping Error: {str(e)}")

@router.post("/documents/verify-forensic-stamp")
async def verify_forensic_stamp(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        filename = file.filename.lower()
        
        author = "UNKNOWN"
        keywords = ""
        comments = ""
        token_data = {}

        if filename.endswith(('.xlsx', '.xls')):
            wb = openpyxl.load_workbook(io.BytesIO(contents), data_only=True)
            props = wb.properties
            author = props.creator or props.lastModifiedBy or "UNKNOWN"
            keywords = props.keywords or ""
            comments = props.description or ""
            
        elif filename.endswith(('.docx', '.doc')):
            doc = Document(io.BytesIO(contents))
            props = doc.core_properties
            author = props.author or props.last_modified_by or "UNKNOWN"
            keywords = props.keywords or ""
            comments = props.comments or ""

        if "TOKEN:" in keywords:
            try:
                raw_token = keywords.split("TOKEN:")[1].split(";")[0]
                token_data = json.loads(base64.b64decode(raw_token).decode("utf-8"))
            except Exception:
                pass

        is_verified = "KMP_CSDMS_AUDIT_STAMP" in keywords or "KMP-STAMP" in comments or bool(token_data)

        return {
            "verified": is_verified,
            "inspected_filename": file.filename,
            "downloaded_by_fnum": token_data.get("fnum", "N/A"),
            "downloaded_by_officer": token_data.get("signature", author),
            "officer_rank": token_data.get("rank", "N/A"),
            "originating_station": token_data.get("station", "N/A"),
            "originating_region": token_data.get("region", "N/A"),
            "export_timestamp": token_data.get("timestamp", "N/A"),
            "forensic_stamp_id": token_data.get("stamp_id", "N/A"),
            "security_classification": "RESTRICTED / ENCRYPTED LAW ENFORCEMENT RECORD" if is_verified else "UNVERIFIED"
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Forensic inspection failed: {str(e)}")

@router.delete("/reports/archive/{doc_id}")
def delete_archive_file(
    doc_id: int, 
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    if current_user.role not in ["SUPER_ADMIN", "ADMIN", "RPC"]:
        raise HTTPException(status_code=403, detail="Command clearance required to delete official records.")
        
    ArchiveModel = get_doc_archive_model()
    TemplateModel = get_template_model()
    GeneralDocModel = get_general_doc_model()
    
    # 🟢 3. Resolve record across Archive, Templates, AND General Documents
    doc_record = db.query(ArchiveModel).filter(ArchiveModel.id == doc_id).first()
    if not doc_record and TemplateModel:
        doc_record = db.query(TemplateModel).filter(TemplateModel.id == doc_id).first()
    if not doc_record and GeneralDocModel:
        doc_record = db.query(GeneralDocModel).filter(GeneralDocModel.id == doc_id).first()

    if not doc_record:
        raise HTTPException(status_code=404, detail="Document not found.")
        
    try:
        file_path = getattr(doc_record, 'file_path', getattr(doc_record, 'url', ''))
        if str(file_path).startswith("http"):
            parsed_url = urllib.parse.urlparse(file_path)
            s3_key = parsed_url.path.lstrip('/')
            try:
                s3_client.delete_object(Bucket=BUCKET_NAME, Key=s3_key)
            except Exception as s3_err:
                print(f"S3 Delete warning: {s3_err}")
            
        db.delete(doc_record)
        db.commit()
        return {"message": "Document successfully deleted from repository and database."}
    except Exception as e:
        print(f"Delete error: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Database Error: {str(e)}")