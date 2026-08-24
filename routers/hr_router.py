import io
import json
import base64
from datetime import datetime
import openpyxl
import pytz
import pyzipper
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import text

from app.database import get_db
from auth import get_current_user

router = APIRouter(prefix="/api/v1/hr", tags=["HR & Establishments"])

def normalize_education_level(educ_str):
    """Normalizes high school levels: keeps uncertified s1-s3 as entered, maps others to UCE or UACE."""
    if not educ_str:
        return "N/A"
    cleaned = str(educ_str).strip().upper()
    
    # Keep uncertified lower secondary classes as entered
    if any(term in cleaned for term in ['S.1', 'S1', 'S.2', 'S2', 'S.3', 'S3', 'SENIOR 1', 'SENIOR 2', 'SENIOR 3']):
        return cleaned
        
    # Map certified levels
    if any(term in cleaned for term in ['UACE', 'A-LEVEL', 'A LEVEL', 'S.6', 'S6', 'SENIOR 6']):
        return "UACE"
    if any(term in cleaned for term in ['UCE', 'O-LEVEL', 'O LEVEL', 'S.4', 'S4', 'SENIOR 4', 'PLE', 'P.7']):
        return "UCE"
        
    return cleaned

@router.get("/export-ledger")
def export_hr_establishments_zip(
    db: Session = Depends(get_db), 
    current_user = Depends(get_current_user)
):
    try:
        # 1. Scope Jurisdiction
        is_global = current_user.role in ['SUPER_ADMIN', 'ADMIN', 'RPC'] or current_user.region in ['KMP HEADQUARTERS', 'POLICE HEADQUARTERS']
        
        nr_query = "SELECT fnum, name, rank, sex, region, station, position, educ_level, status FROM nominal_roll"
        est_query = "SELECT id, region, division, station, personnel_in_station, sub_station, personnel_in_sub_station, post, personnel_in_post, booths, personnel_in_booth, installed_by, location, status, comment, last_updated_by, created_at FROM establishments"
        
        if not is_global:
            nr_query += f" WHERE region = '{current_user.region}'"
            est_query += f" WHERE region = '{current_user.region}'"
            
        nr_records = db.execute(text(nr_query)).fetchall()
        est_records = db.execute(text(est_query)).fetchall()

        # 2. Build Excel File in Memory with Normalized Education Levels
        wb = openpyxl.Workbook()
        ws_nr = wb.active
        ws_nr.title = "Nominal Roll"
        ws_nr.append(["Force Number", "Name", "Rank", "Sex", "Region", "Station", "Position", "Education Level", "Status"])
        for row in nr_records:
            row_list = list(row)
            row_list[7] = normalize_education_level(row_list[7])
            ws_nr.append(row_list)
            
        ws_est = wb.create_sheet(title="establishments")
        ws_est.append(["ID", "Region", "Division", "Station", "Personnel (Station)", "Sub-Station", "Personnel (Sub-Stn)", "Post", "Personnel (Post)", "Booths", "Personnel (Booth)", "Installed By", "Location", "Status", "Comment", "Last Updated By", "Created At"])
        for row in est_records:
            ws_est.append(list(row))

        excel_stream = io.BytesIO()
        wb.save(excel_stream)
        excel_stream.seek(0)

        # 3. Build Formatted Two-Page A4 Landscape Word Document matching UI Structure
        doc = Document()
        
        # Configure A4 Landscape Dimensions and Custom Narrow Margins
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)  # Landscape width
        section.page_height = Inches(8.27)  # Landscape height
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

        # Header Title block
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("KAMPALA METROPOLITAN POLICE - HR & ESTABLISHMENTS LEDGER")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(15, 23, 42)

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run(f"Jurisdiction Scope: {current_user.region} | Generated: {datetime.now().strftime('%Y-%m-%d')}")
        sub_run.font.name = 'Arial'
        sub_run.font.size = Pt(9)
        sub_run.font.color.rgb = RGBColor(100, 116, 139)

        # Page 1 Section: Nominal Roll Table
        h1 = doc.add_paragraph()
        h1_run = h1.add_run("1. Master Personnel Nominal Roll")
        h1_run.font.bold = True
        h1_run.font.size = Pt(10)

        nr_table = doc.add_table(rows=1, cols=8)
        nr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        nr_table.style = 'Table Grid'
        
        hdr_cells = nr_table.rows[0].cells
        headers = ["F/No", "Name", "Rank", "Sex", "Station", "Position", "Educ Level", "Status"]
        for idx, text_val in enumerate(headers):
            hdr_cells[idx].text = text_val
            for p in hdr_cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(8.5)

        for row in nr_records[:150]:
            row_cells = nr_table.add_row().cells
            row_vals = [str(row[0] or ''), str(row[1] or ''), str(row[2] or ''), str(row[3] or ''), str(row[5] or ''), str(row[6] or ''), normalize_education_level(row[7]), str(row[8] or '')]
            for idx, val in enumerate(row_vals):
                row_cells[idx].text = val
                for p in row_cells[idx].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

        # Force Page Break to ensure strict 2-Page landscape layout
        doc.add_page_break()

        # Page 2 Section: Establishments Ledger Table
        h2 = doc.add_paragraph()
        h2_run = h2.add_run("2. Regional Establishments Breakdown")
        h2_run.font.bold = True
        h2_run.font.size = Pt(10)

        est_table = doc.add_table(rows=1, cols=6)
        est_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        est_table.style = 'Table Grid'

        est_hdrs = ["Division", "Station", "Pers (Stn)", "Sub-Station", "Pers (Sub)", "Status"]
        est_hdr_cells = est_table.rows[0].cells
        for idx, text_val in enumerate(est_hdrs):
            est_hdr_cells[idx].text = text_val
            for p in est_hdr_cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(8.5)

        for row in est_records[:100]:
            row_cells = est_table.add_row().cells
            row_vals = [str(row[2] or ''), str(row[3] or ''), str(row[4] or '0'), str(row[5] or '-'), str(row[6] or '0'), str(row[13] or 'OPERATIONAL')]
            for idx, val in enumerate(row_vals):
                row_cells[idx].text = val
                for p in row_cells[idx].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)

        # 4. Encrypt inside a ZIP using Force Number
        eat_tz = pytz.timezone("Africa/Nairobi")
        eat_time = datetime.now(eat_tz).replace(tzinfo=None)
        
        officer_fnum = (current_user.fnum or "UNKNOWN").strip().upper()
        zip_stream = io.BytesIO()
        zip_password = str(current_user.fnum).strip().encode('utf-8')

        with pyzipper.AESZipFile(zip_stream, 'w', compression=pyzipper.ZIP_DEFLATED, encryption=pyzipper.WZ_AES) as zf:
            zf.setpassword(zip_password)
            excel_filename = f"{officer_fnum.replace('/', '_')}_HR_Ledger_{eat_time.strftime('%Y%m%d')}.xlsx"
            word_filename = f"{officer_fnum.replace('/', '_')}_HR_Ledger_Report_{eat_time.strftime('%Y%m%d')}.docx"
            
            zf.writestr(excel_filename, excel_stream.getvalue())
            zf.writestr(word_filename, doc_stream.getvalue())

        zip_stream.seek(0)

        # 5. Send ZIP to Frontend
        zip_filename = f"SECURE_HR_LEDGER_{eat_time.strftime('%Y%m%d')}.zip"
        headers = {
            'Content-Disposition': f'attachment; filename="{zip_filename}"',
            'Access-Control-Expose-Headers': 'Content-Disposition'
        }
        
        return StreamingResponse(
            zip_stream, 
            media_type="application/zip",
            headers=headers
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export Error: {str(e)}")